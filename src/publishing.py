import docx


class DocWriter:

    def __init__(self) -> None:
        self.doc = docx.Document()

    def write_doc(self, book, chapter_dict, title):

        self.doc.add_heading(title, 0)

        for chapter, paragraphs_list in book.items():

            description = chapter_dict[chapter]
            chapter_name = "{}: {}".format(chapter.strip(), description.strip())

            self.doc.add_page_break()
            self.doc.add_heading(chapter_name, 1)

            text = "\n\n".join(paragraphs_list)
            self.doc.add_paragraph(text)

        self.doc.save("./docs/book_1.docx")
