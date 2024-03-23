import re
from docx import Document


def modify(docx_file):

    doc = Document(docx_file)

    # Define regular expression patterns for "In summary" and "In conclusion"
    summary_pattern = re.compile(r"^\s*In\s+summary", re.IGNORECASE)
    conclusion_pattern = re.compile(r"^\s*In\s+conclusion", re.IGNORECASE)

    # Create a new list to store filtered paragraphs
    filtered_paragraphs = []

    # Iterate over each paragraph in the original document
    for paragraph in doc.paragraphs:
        # Check if the paragraph starts with "In summary" or "In conclusion"
        if not summary_pattern.match(
            paragraph.text.strip()
        ) and not conclusion_pattern.match(paragraph.text.strip()):
            filtered_paragraphs.append(paragraph)

    # Clear the original document
    for paragraph in doc.paragraphs:
        paragraph.clear()

    # Add the filtered paragraphs back to the original document
    for paragraph in filtered_paragraphs:
        doc.add_paragraph(paragraph.text)

    # Save the modified document
    modified_docx_file = docx_file.replace(".docx", "_modified.docx")
    doc.save(modified_docx_file)

    print(
        f"Summary and conclusion paragraphs removed. Modified document saved as: {modified_docx_file}"
    )


modify("./docs/book_1_GPT3.docx")
